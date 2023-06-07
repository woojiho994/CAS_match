# -*- coding: utf-8 -*-
"""
Created on Sun Jun  4 15:55:12 2023

@author: wooji
"""
import re
import streamlit as st
import numpy as np
import pandas as pd
import streamlit as st
import pandas as pd
from io import StringIO
import pdfplumber
import xlrd
import re
import pdfplumber
import pandas as pd
import time
import os
import numpy as np
#import win32com
#from win32com.client import Dispatch
import docx2pdf
import docx
#import win32com.client as wc
#import win32com.client as win32
import pytesseract
from PIL import  Image
import os
from pdf2image import convert_from_path,convert_from_bytes
from io import BytesIO

import openpyxl
import base64
st.title("MSDS报告CAS号提取程序")
#%%
from docx import Document

def get_tables(docx_path):
    docStr = Document(docx_path)
    numTables = docStr.tables
    my_list = []
    for table in numTables:
        row_count = len(table.rows)       
        col_count = len(table.columns)
        for i in range(row_count):
            row = table.rows[i].cells
            for j in range(col_count):     
                    content = row[j].text  
                    my_list.append(content)
    my_list = ';'.join(my_list).strip('')
    return my_list


def get_paragraphs(docx_path):
    #打开word文档
    document = Document(docx_path)  
    #获取所有段落
    all_paragraphs = document.paragraphs    
    paragraph_texts = []
    # 循环读取列表
    for paragraph in all_paragraphs:
        paragraph_texts.append(paragraph.text)
    paragraph_texts = ';'.join(paragraph_texts).strip('')
    return paragraph_texts



#%% 函数二、打开pdf文件,输出每一页pdf中的所有文字
def openpdf(path):
    with pdfplumber.open(path) as pdf:
        # pdf = pdfplumber.open(path)
        item = []
        for page in pdf.pages:
            text = page.extract_text()
            item.append(text)
        # item = [''.join(i) for i in item]
        item = ';'.join(item).strip('')
    return item

#%% 函数三、将目标CAS号，和pdf中的内容进行比对。返回什么？
def extract(text,cas):
    pattern = re.compile(cas,re.S)
    r_list = pattern.findall(text)
    return r_list
#%%
# data = pd.DataFrame(columns=['CAS','名称','匹配结果','备注'])
st.write('使用说明')
st.caption('支持解析的格式：.pdf(扫描版或非扫描版均支持)和.docx。可将MSDS文件夹直接拖拽到下方上传区域')
st.write('excel输出内容详解')
st.caption('第一列为文件名称,所有上传的文件均会显示在第一列，即便该文件格式不支持提取')
st.caption('第二列为文件中提取的CAS号,若为空则表明未提取到')
st.caption('第三列为化学物质名称,仅支持显示与清单匹配成功的化学物质的名称')
st.caption('第四列为匹配结果,共3种结果：3960种、优评优控、重点管控')
st.caption('第五列为备注,共3种结果：1、不支持该格式文件，请手动查看：说明此类文件不支持解析，请手动查看；2、图片pdf，建议人工复核：说明该pdf为图片，提取正确率较低，视情况可进行人工复核；3、未检测到CAS，请手动检查：说明在该文件中未检测到CAS，请人工确认')

st.caption('提取速度：提取一个电子pdf大约耗时4s，一个扫描版pdf大约耗时10~20s。具体速度由pdf的页数决定')
st.divider()
uploaded_file = st.file_uploader("请上传MSDS报告，可直接往里拖拽文件夹",accept_multiple_files=True)
@st.cache_data
def main(uploaded_file):
    data = pd.DataFrame(columns=['CAS','名称','匹配结果','备注'])

    begin = time.time()
    # openpdf(uploaded_file)
    cas = r'[0-9]+-[0-9][0-9]-[0-9][^0-9]' 
    # st.write(extract(openpdf(uploaded_file),cas))
    for file in range(len(uploaded_file)):
        if uploaded_file[file].name[-4:] == 'docx':
            text =  get_paragraphs(uploaded_file[file])
            # text(get_tables(uploaded_file[file]))
            # text = ';'.join(text).strip('')                       
        elif uploaded_file[file].name[-3:] == 'pdf' or uploaded_file[file].name[-3:] == 'PDF':  
            text = openpdf(uploaded_file[file])
        else:
            cas_set = pd.DataFrame({'备注':{uploaded_file[file].name:'不支持该格式文件，请手动查看'}})
            data = pd.concat([data,cas_set],axis=0) 
            continue
        cas_extract = extract(text,cas)
        if cas_extract != []:
            for item in range(len(cas_extract)):
                cas_iso = cas_extract[item]
                cas_iso = cas_iso[0:len(cas_iso)-1]
                cas_set = pd.DataFrame({'CAS':{uploaded_file[file].name:cas_iso}})
                data = pd.concat([data,cas_set],axis=0) 
        #提取docx表格内的内容
        elif uploaded_file[file].name[-4:] == 'docx':   
            text = get_tables(uploaded_file[file])
            # text = ';'.join(text).strip('') 
            cas_extract = extract(text,cas)   
            if cas_extract != []:
                for item in range(len(cas_extract)):
                    cas_iso = cas_extract[item]
                    cas_iso = cas_iso[0:len(cas_iso)-1]
                    cas_set = pd.DataFrame({'CAS':{uploaded_file[file].name:cas_iso}})
                    data = pd.concat([data,cas_set],axis=0)
            else:
                cas_set = pd.DataFrame({'备注':{uploaded_file[file].name:'未检测到CAS，请手动检查'}})
                data = pd.concat([data,cas_set],axis=0)
        else:
            pages = convert_from_bytes(uploaded_file[file].getvalue()) # 上传的内容是什么？
            text = []
            for i,page in enumerate(pages):
                 buf = BytesIO()
                 page.save(buf,format="JPEG")
                 buf.seek(0)
                 img_page=Image.open(buf)
                 # st.write('here')
                 txt=pytesseract.image_to_string(img_page,lang='chi_sim')
                 text.append(txt)  
            text = ';'.join(text).strip('')
            cas_extract = extract(text,cas)
            if cas_extract != []:
                cas_extract = extract(text,cas)
                for item in range(len(cas_extract)):
                    cas_iso = cas_extract[item]
                    cas_iso = cas_iso[0:len(cas_iso)-1]
                    print(cas_iso)
                    # cas_set = pd.Series({uploaded_file[file].name:cas_iso+'图片pdf，请手动检查'})   #在这里加备注提示是扫描版pdf
                    #用dataframe承载
                    cas_set = pd.DataFrame({'CAS':{uploaded_file[file].name:cas_iso},'备注':{uploaded_file[file].name:'图片pdf，建议人工复核'}})
                    data = pd.concat([data,cas_set],axis=0)
            else:
                cas_set = pd.DataFrame({'备注':{uploaded_file[file].name:'未检测到CAS，请手动检查'}})
                data = pd.concat([data,cas_set],axis=0)
    
    # st.write(uploaded_file)
    # convert_from_bytes(open('/home/belval/example.pdf','rb').read())
    
    

#%%数据整理
    data_reset_index = data.reset_index(drop=False)  
    #修改列名
    data_rename = data_reset_index.rename(columns={'index':'MSDS文件名称'})    
    #去除重复行             
    data_output = data_rename.drop_duplicates()  #subset='pdf名称'可以查看是不是所有文件都包含在表格里
    # target_data_base = pd.read_excel('C:/Users/wooji/Nutstore/1/Jiho华南所/鉴定中心-工作/MSDS/102-104物质清单.xlsx',sheet_name='基102-3960种',index_col=0)
    # # target_data_pri = pd.read_excel('C:/Users/wooji/Nutstore/1/Jiho华南所/鉴定中心-工作/MSDS/物质清单.xlsx',sheet_name='优评优控',index_col=0)
    # # target_data_key = pd.read_excel('C:/Users/wooji/Nutstore/1/Jiho华南所/鉴定中心-工作/MSDS/物质清单.xlsx',sheet_name='重点管控',index_col=0)
    # target_cas_base = target_data_base[['CAS','名称']]
    # target_cas_pri = target_data_pri[['CAS','名称']]
    # target_cas_key = target_data_key[['CAS','名称']]
    # target_cas_base = target_cas_base.reset_index(drop=True)
    # target_cas_pri = target_cas_pri.reset_index(drop=True)
    # target_cas_key = target_cas_key.reset_index(drop=True)
    target_data = pd.read_excel('C:/Users/wooji/Nutstore/1/Jiho华南所/鉴定中心-工作/MSDS/物质清单.xlsx',sheet_name='总表',index_col=0)
    target_cas = target_data[['CAS','名称','清单']]
    target_cas = target_cas.reset_index(drop=True)

    
#%%
    for row in data_output.index:
        # print(data_output.loc[row]['CAS号提取'])
        for b in target_cas.index:
            if data_output.loc[row]['CAS'] == target_cas.loc[b]['CAS']:
                data_output.loc[row]['匹配结果'] =target_cas.loc[b]['清单']
                data_output.loc[row]['名称'] = target_cas.loc[b]['名称']
    
 
    data_final = data_output
    # [['pdf名称','匹配结果','CAS号提取','名称','备注']]
    end = time.time()
    run_time = end - begin
    st.write('运行耗时：'+ str(round(run_time,2))+'秒')
    return data_final


if uploaded_file == []:
    st.stop()
else:  
    data_final = main(uploaded_file)
    data_final
    data_final.to_excel('resuls.xlsx')
    wb2 = openpyxl.load_workbook('resuls.xlsx')
    wb2.save('results.xlsx')#注意！文件此时保存在内存中且为字节格式文件
    data=open('results.xlsx','rb').read()#以只读模式读取且读取为二进制文件
    b64 = base64.b64encode(data).decode('UTF-8')#解码并加密为base64
    excel_name = st.text_input(':blue[请输入本次导入的文件所属企业名称，若为空则导出的excel默认取名为myresult.xlsx]')
    st.warning('建议示例：广西xx企业-原辅料 or 广西xx企业-产品  ------- 输入完请按回车 ', icon="🚨")
if excel_name:
    excel_name = excel_name + '.xlsx'
    href = f'<a href="data:file/data;base64,{b64}" download={excel_name}>导出excel</a>'#定义下载链接，默认的下载文件名是myresults.xlsx
    st.markdown(href, unsafe_allow_html=True)#输出到浏览器
    wb2.close()
else:
    href = f'<a href="data:file/data;base64,{b64}" download=myresult.xlsx>导出excel</a>'#定义下载链接，默认的下载文件名是myresults.xlsx
    st.markdown(href, unsafe_allow_html=True)#输出到浏览器
    wb2.close()
        

st.subheader('!!!单次使用完请刷新页面后再上传新的文件')


# else:
#     excel_name = excel_name + '.xlsx'
#     href = f'<a href="data:file/data;base64,{b64}" download={excel_name}>Download xlsx file</a>'#定义下载链接，默认的下载文件名是myresults.xlsx
#     st.markdown(href, unsafe_allow_html=True)#输出到浏览器
#     wb2.close()




####直接写识别图片的代码
# stringio = StringIO(uploaded_file[file].getvalue().decode("utf-8"))
# st.write(stringio)   ##这句是对的
# bytes_data = uploaded_file[file].read()
# st.write(bytes_data)
# st.write(uploaded_file[file])
# st.write(bytes_data)
# =============================================================================
#             ####
#             stringio = StringIO(uploaded_file[file].getvalue().decode("utf-8"))
#             st.write(stringio)
#             # To read file as string:
#             string_data = stringio.read()
#             st.write(string_data)
#             ###
# =============================================================================





